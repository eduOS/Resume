#!/usr/bin/env python3
# Simple Markdown to DOCX converter for this resume.
# Handles headings (#, ##, ###), bullet lists (- or •), and plain paragraphs.
# Usage: python3 md_to_docx.py input.md output.docx

import sys
import re
from docx import Document

def convert(md_path, docx_path):
    doc = Document()
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    for raw in lines:
        line = raw.rstrip('\n')
        if line.strip().startswith('```'):
            in_code_block = not in_code_block
            continue
        if in_code_block:
            p = doc.add_paragraph(line)
            p.style = 'Intense Quote'
            continue
        s = line.strip()
        if not s:
            # blank line -> paragraph break
            doc.add_paragraph('')
            continue
        # headings
        if s.startswith('# '):
            doc.add_heading(s[2:].strip(), level=1)
            continue
        if s.startswith('## '):
            doc.add_heading(s[3:].strip(), level=2)
            continue
        if s.startswith('### '):
            doc.add_heading(s[4:].strip(), level=3)
            continue
        # bullets
        m = re.match(r'^[-•\*]\s+(.*)', s)
        if m:
            p = doc.add_paragraph(m.group(1).strip())
            try:
                p.style = 'List Bullet'
            except Exception:
                # fallback: prepend a bullet char
                p.text = '• ' + p.text
            continue
        # bold markers **text** -> remove markers (simple)
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', s)
        doc.add_paragraph(text)

    doc.save(docx_path)
    print('Saved:', docx_path)

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print('Usage: md_to_docx.py input.md output.docx')
        sys.exit(1)
    convert(sys.argv[1], sys.argv[2])
